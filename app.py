from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify, Response
import os
from datetime import datetime
from docx import Document
import mimetypes

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'

# Configuration
UPLOAD_FOLDER = 'documents'
ALLOWED_EXTENSIONS = {'docx', 'doc'}

# Create documents folder if it doesn't exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_documents():
    """Get list of all Word documents in the upload folder"""
    documents = []
    try:
        for filename in os.listdir(UPLOAD_FOLDER):
            if allowed_file(filename):
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                stat = os.stat(filepath)
                documents.append({
                    'name': filename,
                    'size': round(stat.st_size / 1024, 2),
                    'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                    'path': filepath
                })
    except Exception as e:
        print(f"Error reading documents: {e}")
    return sorted(documents, key=lambda x: x['modified'], reverse=True)

@app.route('/')
def index():
    """Main page showing all documents"""
    documents = get_documents()
    return render_template('index.html', documents=documents)

@app.route('/create', methods=['GET', 'POST'])
def create_document():
    """Create a new Word document"""
    if request.method == 'POST':
        try:
            filename = request.form['filename'].strip()
            title = request.form.get('title', '').strip()
            content = request.form.get('content', '').strip()
            
            if not filename:
                flash('Filename is required', 'error')
                return render_template('create.html')
            
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            # Handle duplicate filenames
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            counter = 1
            base_name, ext = os.path.splitext(filename)
            while os.path.exists(filepath):
                filename = f"{base_name}_{counter}{ext}"
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                counter += 1
            
            # Create new document
            doc = Document()
            
            # Add title
            if title:
                doc.add_heading(title, 0)
            
            # Add content
            if content:
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            
            # Save document
            doc.save(filepath)
            
            flash(f'Document "{filename}" created successfully!', 'success')
            return redirect(url_for('index'))
            
        except Exception as e:
            flash(f'Error creating document: {str(e)}', 'error')
    
    return render_template('create.html')

@app.route('/upload', methods=['GET', 'POST'])
def upload_document():
    """Upload existing Word document"""
    if request.method == 'POST':
        try:
            if 'file' not in request.files:
                flash('No file selected', 'error')
                return redirect(request.url)
            
            file = request.files['file']
            if file.filename == '':
                flash('No file selected', 'error')
                return redirect(request.url)
            
            if file and allowed_file(file.filename):
                filename = file.filename
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                
                # Handle duplicate filenames
                counter = 1
                base_name, ext = os.path.splitext(filename)
                while os.path.exists(filepath):
                    filename = f"{base_name}_{counter}{ext}"
                    filepath = os.path.join(UPLOAD_FOLDER, filename)
                    counter += 1
                
                file.save(filepath)
                flash(f'Document "{filename}" uploaded successfully!', 'success')
                return redirect(url_for('index'))
            else:
                flash('Invalid file type. Please upload .docx or .doc files only.', 'error')
                
        except Exception as e:
            flash(f'Error uploading document: {str(e)}', 'error')
    
    return render_template('upload.html')

@app.route('/edit/<filename>')
def edit_document(filename):
    """Download document for editing on user's computer"""
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(filepath):
        flash('Document not found', 'error')
        return redirect(url_for('index'))
    
    # Set the proper MIME type to trigger Word application
    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    if filename.endswith('.doc'):
        mimetype = 'application/msword'
    
    try:
        # Send file with headers that encourage opening in Word
        return send_file(
            filepath,
            mimetype=mimetype,
            as_attachment=False,  # This makes it open instead of download
            download_name=filename
        )
    except Exception as e:
        flash(f'Error opening document: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/webedit/<filename>', methods=['GET', 'POST'])
def web_edit_document(filename):
    """Web-based document editing"""
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(filepath):
        flash('Document not found', 'error')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        try:
            # Save the edited content
            content = request.form.get('content', '').strip()
            title = request.form.get('title', '').strip()
            
            # Create new document with updated content
            doc = Document()
            
            if title:
                doc.add_heading(title, 0)
            
            # Split content by lines and add as paragraphs
            if content:
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            
            doc.save(filepath)
            flash(f'Document "{filename}" updated successfully!', 'success')
            return redirect(url_for('index'))
            
        except Exception as e:
            flash(f'Error saving document: {str(e)}', 'error')
    
    # Read current content for editing
    try:
        doc = Document(filepath)
        content = []
        title = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') and not title:
                title = paragraph.text
            elif paragraph.text.strip():
                content.append(paragraph.text)
        
        content_text = '\n'.join(content)
        return render_template('webedit.html', filename=filename, content=content_text, title=title)
        
    except Exception as e:
        flash(f'Error reading document: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download_document(filename):
    """Download document"""
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(filepath):
        flash('Document not found', 'error')
        return redirect(url_for('index'))
    
    return send_file(filepath, as_attachment=True)

@app.route('/delete/<filename>')
def delete_document(filename):
    """Delete document"""
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(filepath):
        flash('Document not found', 'error')
        return redirect(url_for('index'))
    
    try:
        os.remove(filepath)
        flash(f'Document "{filename}" deleted successfully!', 'success')
    except Exception as e:
        flash(f'Error deleting document: {str(e)}', 'error')
    
    return redirect(url_for('index'))

@app.route('/preview/<filename>')
def preview_document(filename):
    """Preview document content"""
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(filepath):
        flash('Document not found', 'error')
        return redirect(url_for('index'))
    
    try:
        doc = Document(filepath)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        return render_template('preview.html', filename=filename, content=content)
    except Exception as e:
        flash(f'Error reading document: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/api/documents')
def api_documents():
    """API endpoint to get documents list"""
    return jsonify(get_documents())

# Route for handling file uploads after editing
@app.route('/upload_edited/<filename>', methods=['POST'])
def upload_edited_document(filename):
    """Handle upload of edited document"""
    if 'file' not in request.files:
        flash('No file uploaded', 'error')
        return redirect(url_for('index'))
    
    file = request.files['file']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    
    if file and allowed_file(file.filename):
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        flash(f'Document "{filename}" updated successfully!', 'success')
    else:
        flash('Invalid file type', 'error')
    
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
